import telebot
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os
from io import BytesIO

TOKEN = "8588192735:AAFS1MUM5LINj6hWIpe8f6xgfKs9UyosC_c"
bot = telebot.TeleBot(TOKEN)

# Словари для хранения состояния пользователя
user_states = {}
user_profession_matches = {}
user_selected_profession = {}
user_selected_hazards = {}
user_hazard_matches = {}
professions_df = None
hazards_df = None


def load_professions_df():
    """Загружает DataFrame профессий один раз"""
    global professions_df
    if professions_df is None:
        try:
            professions_df = pd.read_excel('professions.xlsx', header=None)
        except Exception as e:
            print(f"Ошибка при загрузке профессий: {e}")
            return None
    return professions_df


def load_hazards_df():
    """Загружает DataFrame опасностей один раз"""
    global hazards_df
    if hazards_df is None:
        try:
            hazards_df = pd.read_excel('iden_haz.xlsx', header=None)
        except Exception as e:
            print(f"Ошибка при загрузке опасностей: {e}")
            return None
    return hazards_df


def parse_norm_string(norm_str):
    """Парсит строку норм выдачи и извлекает количество и единицу измерения"""
    if pd.isna(norm_str):
        return None, None

    norm_str = str(norm_str).strip()

    if "износа" in norm_str.lower():
        return "до износа", None

    match = re.match(r'(\d+)\s*([а-яА-Я]+)', norm_str)
    if match:
        quantity = int(match.group(1))
        unit = match.group(2).strip()

        unit_map = {
            'пара': 'пары',
            'пары': 'пары',
            'шт': 'штуки',
            'штука': 'штуки',
            'штуки': 'штуки',
            'комплект': 'комплекты',
            'комплекты': 'комплекты',
            'мл': 'мл'
        }

        normalized_unit = unit_map.get(unit.lower(), unit)
        return quantity, normalized_unit

    return None, None


def extract_years_info(norm_str):
    """Извлекает информацию о периоде"""
    if pd.isna(norm_str):
        return ""

    norm_str = str(norm_str).lower()

    if "на 2 года" in norm_str:
        return " на 2 года"
    elif "на 3 года" in norm_str:
        return " на 3 года"
    elif "на 5 лет" in norm_str:
        return " на 5 лет"
    elif "на 1,5 года" in norm_str:
        return " на 1,5 года"

    return ""


def find_profession_by_code(professions_df, code):
    """Ищет профессию по коду"""
    try:
        code = int(code)
    except:
        return None, None, None

    matches = professions_df[professions_df.iloc[:, 0] == code]

    if matches.empty:
        return None, None, None

    start_idx = matches.index[0]
    profession_name = professions_df.iloc[start_idx, 1]

    return code, profession_name, start_idx


def find_professions_by_name(professions_df, search_name, limit=10):
    """Ищет профессии по названию с ограничением результатов"""
    search_name_lower = search_name.lower().strip()
    matches = []

    for idx, row in professions_df.iterrows():
        if pd.notna(row.iloc[0]) and isinstance(row.iloc[0], (int, float)):
            code = int(row.iloc[0])
            profession_name = row.iloc[1]

            if pd.notna(profession_name):
                profession_name_str = str(profession_name).strip()
                if search_name_lower in profession_name_str.lower():
                    matches.append({
                        'code': code,
                        'name': profession_name_str,
                        'idx': idx
                    })

                if len(matches) >= limit:
                    break

    return matches


def find_hazards_by_keyword(hazards_df, keyword):
    """Ищет опасности по ключевому слову в столбце A (индекс 0)"""
    keyword_lower = keyword.lower().strip()
    matches = []
    seen_hazards = set()

    for idx, row in hazards_df.iterrows():
        hazard_name = row.iloc[0]

        if pd.notna(hazard_name):
            hazard_name_str = str(hazard_name).strip()
            if keyword_lower in hazard_name_str.lower() and hazard_name_str not in seen_hazards:
                matches.append({
                    'name': hazard_name_str,
                    'idx': idx
                })
                seen_hazards.add(hazard_name_str)

    return matches


def extract_siz_from_hazard(hazards_df, hazard_idx):
    """Извлекает СИЗ из строки опасности в файле iden_haz.xlsx

    Структура файла:
    Столбец A (0) - Название опасности
    Столбец C (2) - Тип/группа СИЗ
    Столбец D (3) - Конструкция СИЗ (наименование)
    Столбец E (4) - Нормы выдачи
    """
    siz_list = []
    row = hazards_df.iloc[hazard_idx]

    # Начинаем со следующей строки и собираем связанные СИЗ
    # пока не встретим новую опасность (новую заполненную ячейку в столбце A)
    idx = hazard_idx + 1

    while idx < len(hazards_df):
        current_row = hazards_df.iloc[idx]

        # Если встретили новую опасность в столбце A, останавливаемся
        if pd.notna(current_row.iloc[0]) and str(current_row.iloc[0]).strip():
            # Проверяем, это ли опасность (начинается ли строка с нового номера)
            if current_row.iloc[0] != "или":
                break

        # Извлекаем наименование СИЗ из столбца D (индекс 3)
        siz_name = current_row.iloc[3] if len(current_row) > 3 else None

        # Извлекаем норму выдачи из столбца E (индекс 4)
        norm_str = current_row.iloc[4] if len(current_row) > 4 else None

        if pd.notna(siz_name) and str(siz_name).strip():
            siz_name_str = str(siz_name).strip()

            # Пропускаем служебные строки
            if siz_name_str.lower() not in ['или', 'и/или', 'in', 'or']:
                quantity, unit = parse_norm_string(norm_str)
                years_info = extract_years_info(norm_str)

                if quantity:
                    quantity_str = f"{quantity}{years_info}"
                else:
                    quantity_str = str(norm_str) if pd.notna(norm_str) and str(norm_str).strip() else "до износа"

                siz_list.append({
                    'name': siz_name_str,
                    'unit': unit if unit else 'штуки',
                    'quantity': quantity_str,
                    'norm': '№767н'
                })

        idx += 1

    return siz_list


def extract_siz_data(professions_df, profession_code, start_idx, hazards_list=None):
    """Извлекает данные о СИЗ для конкретной профессии и опасностей"""
    siz_dict = {}

    # Извлекаем СИЗ из профессии
    for idx in range(start_idx + 1, len(professions_df)):
        row = professions_df.iloc[idx]

        if pd.notna(row.iloc[0]) and isinstance(row.iloc[0], (int, float)) and not pd.isna(row.iloc[0]):
            break

        siz_name = row.iloc[3] if len(row) > 3 else None
        norm_str = row.iloc[4] if len(row) > 4 else None

        if pd.notna(siz_name) and str(siz_name).strip():
            quantity, unit = parse_norm_string(norm_str)
            years_info = extract_years_info(norm_str)

            if quantity:
                quantity_str = f"{quantity}{years_info}"
            else:
                quantity_str = str(norm_str) if pd.notna(norm_str) else "до износа"

            siz_name_str = str(siz_name).strip()

            if siz_name_str not in siz_dict:
                siz_dict[siz_name_str] = {
                    'name': siz_name_str,
                    'unit': unit if unit else 'штуки',
                    'quantity': quantity_str,
                    'norm': '№767н'
                }

    # Извлекаем СИЗ из опасностей
    if hazards_list:
        hazards_df = load_hazards_df()

        for hazard_name in hazards_list:
            # Ищем строку с этой опасностью
            for idx, row in hazards_df.iterrows():
                current_hazard = row.iloc[0]
                if pd.notna(current_hazard):
                    current_hazard_str = str(current_hazard).strip()
                    if current_hazard_str == hazard_name:
                        # Нашли опасность, извлекаем связанные СИЗ
                        hazard_siz_list = extract_siz_from_hazard(hazards_df, idx)

                        for siz in hazard_siz_list:
                            if siz['name'] not in siz_dict:
                                siz_dict[siz['name']] = siz
                            else:
                                # Если СИЗ уже есть, обновляем количество
                                existing = siz_dict[siz['name']]
                                # Объединяем количества через запятую если они разные
                                if existing['quantity'] != siz['quantity']:
                                    existing['quantity'] = f"{existing['quantity']}, {siz['quantity']}"
                        break

    return list(siz_dict.values())


def set_cell_border(cell, **kwargs):
    """Устанавливает границы для ячейки таблицы"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for edge in ('top', 'left', 'bottom', 'right'):
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), 'single')
        edge_el.set(qn('w:sz'), '12')
        edge_el.set(qn('w:space'), '0')
        edge_el.set(qn('w:color'), '000000')
        tcBorders.append(edge_el)

    tcPr.append(tcBorders)


def find_siz_table(doc):
    """Находит таблицу СИЗ в документе Word"""
    for idx, table in enumerate(doc.tables):
        if len(table.rows) > 0:
            first_cell = table.rows[0].cells[0].text.strip()
            if "Наименование СИЗ" in first_cell:
                return table, idx

    return None, None


def add_table_row(table, siz_data):
    """Добавляет новую строку в таблицу"""
    new_row = table.add_row()

    if len(table.rows) > 2:
        source_row = table.rows[-2]
        for source_cell, target_cell in zip(source_row.cells, new_row.cells):
            set_cell_border(target_cell)
    else:
        for cell in new_row.cells:
            set_cell_border(cell)

    return new_row


def fill_siz_table(doc, siz_data):
    """Заполняет таблицу СИЗ данными"""
    table, table_idx = find_siz_table(doc)

    if not table:
        return False

    header_rows = 1
    existing_data_rows = len(table.rows) - header_rows
    required_rows = len(siz_data)

    if required_rows > existing_data_rows:
        rows_to_add = required_rows - existing_data_rows
        for i in range(rows_to_add):
            add_table_row(table, siz_data)

    for idx, siz in enumerate(siz_data):
        row_idx = header_rows + idx
        row = table.rows[row_idx]

        if len(row.cells) >= 4:
            for cell in row.cells[:4]:
                cell.text = ""

            row.cells[0].text = siz['name']
            row.cells[1].text = siz['norm']
            row.cells[2].text = siz['unit']
            row.cells[3].text = siz['quantity']

    return True


def generate_siz_document_from_match(match, hazards_list=None):
    """Генерирует документ из объекта профессии"""
    try:
        professions_df = load_professions_df()
        if professions_df is None:
            return None, "Ошибка при загрузке файла профессий"

        siz_data = extract_siz_data(professions_df, match['code'], match['idx'], hazards_list)

        if not siz_data:
            return None, f"❌ Не найдены СИЗ для профессии '{match['name']}'"

        try:
            doc = Document('personal_anketa.docx')
        except Exception as e:
            return None, f"Ошибка при открытии шаблона: {str(e)}"

        if not fill_siz_table(doc, siz_data):
            return None, "Ошибка при заполнении таблицы"

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return output, None

    except Exception as e:
        return None, f"Ошибка при генерации документа: {str(e)}"


@bot.message_handler(commands=['start'])
def send_welcome(message):
    welcome_text = """Привет, я твой помощник в сфере Техносферной безопасности. Я могу:

• Рассчитать нормы выдачи СИЗ
• Помочь вспомнить терминологию
• Найти соответствующий закон

Выбери, что тебя интересует:"""

    markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    btn1 = telebot.types.KeyboardButton("📋 Расчет СИЗ")
    btn2 = telebot.types.KeyboardButton("📚 Терминология")
    btn3 = telebot.types.KeyboardButton("⚖️ Законодательство")

    markup.add(btn1, btn2, btn3)
    bot.send_message(message.chat.id, welcome_text, reply_markup=markup)


@bot.message_handler(commands=['help'])
def send_help(message):
    help_text = """Доступные команды:
/start - Начать работу с ботом
/help - Справка

Основные функции:
📋 Расчет СИЗ - Рассчитать нормы выдачи средств индивидуальной защиты
📚 Терминология - Справочник терминов в области техносферной безопасности
⚖️ Законодательство - Информация о соответствующих законах и нормативах"""

    markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    btn_start = telebot.types.KeyboardButton("Начать")
    btn_help = telebot.types.KeyboardButton("Помощь")

    markup.add(btn_start, btn_help)
    bot.send_message(message.chat.id, help_text, reply_markup=markup)


@bot.message_handler(func=lambda message: message.text == "Начать")
def on_start_button(message):
    send_welcome(message)


@bot.message_handler(func=lambda message: message.text == "Помощь")
def on_help_button(message):
    send_help(message)


@bot.message_handler(func=lambda message: message.text == "📋 Расчет СИЗ")
def calculate_siz(message):
    chat_id = message.chat.id
    user_states[chat_id] = "waiting_profession"

    request_text = """🔍 Введи название или код профессии:

Начни писать, и я предложу похожие специальности.

Примеры поиска:
• Авиационный
• Сторож
• 4732
• Водитель"""

    markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    markup.add(telebot.types.KeyboardButton("◀️ Назад"))

    msg = bot.send_message(chat_id, request_text, reply_markup=markup)
    bot.register_next_step_handler(msg, process_profession_input)


def process_profession_input(message):
    chat_id = message.chat.id

    if message.text == "◀️ Назад":
        user_states.pop(chat_id, None)
        user_profession_matches.pop(chat_id, None)
        send_welcome(message)
        return

    profession_query = message.text.strip()

    if not profession_query or len(profession_query) < 2:
        markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
        markup.add(telebot.types.KeyboardButton("◀️ Назад"))
        msg = bot.send_message(chat_id, "❌ Введите минимум 2 символа для поиска\n\nПопробуйте еще раз:",
                               reply_markup=markup)
        bot.register_next_step_handler(msg, process_profession_input)
        return

    professions_df = load_professions_df()
    if professions_df is None:
        bot.send_message(chat_id, "❌ Ошибка при загрузке файла профессий")
        user_states.pop(chat_id, None)
        return

    code, profession_name, start_idx = find_profession_by_code(professions_df, profession_query)

    if code is not None:
        selected_profession = {
            'code': code,
            'name': profession_name,
            'idx': start_idx
        }
        user_selected_profession[chat_id] = selected_profession
        ask_for_additional_hazards(message, selected_profession)
        return

    matches = find_professions_by_name(professions_df, profession_query, limit=10)

    if not matches:
        markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
        markup.add(telebot.types.KeyboardButton("◀️ Назад"))
        msg = bot.send_message(chat_id,
                               f"❌ Профессия '{profession_query}' не найдена\n\nПопробуйте еще раз с другим названием:",
                               reply_markup=markup)
        bot.register_next_step_handler(msg, process_profession_input)
        return

    if len(matches) == 1:
        user_selected_profession[chat_id] = matches[0]
        ask_for_additional_hazards(message, matches[0])
        return

    user_profession_matches[chat_id] = matches
    user_states[chat_id] = "waiting_profession_choice"

    message_text = "🔍 Найдено несколько похожих профессий:\n\n"

    markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)

    for i, match in enumerate(matches, 1):
        message_text += f"{i}. {match['name']} (№ {match['code']})\n"
        markup.add(telebot.types.KeyboardButton(str(i)))

    message_text += "\n👇 Выбери номер нужной профессии:"

    markup.add(telebot.types.KeyboardButton("◀️ Назад"))
    msg = bot.send_message(chat_id, message_text, reply_markup=markup)
    bot.register_next_step_handler(msg, process_profession_choice)


def process_profession_choice(message):
    chat_id = message.chat.id

    if message.text == "◀️ Назад":
        user_states[chat_id] = "waiting_profession"
        user_profession_matches.pop(chat_id, None)

        request_text = """🔍 Введи название или код профессии:

Начни писать, и я предложу похожие специальности."""

        markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
        markup.add(telebot.types.KeyboardButton("◀️ Назад"))

        msg = bot.send_message(chat_id, request_text, reply_markup=markup)
        bot.register_next_step_handler(msg, process_profession_input)
        return

    choice_text = message.text.strip()

    try:
        choice_num = int(choice_text)
    except ValueError:
        markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)

        matches = user_profession_matches.get(chat_id, [])
        for i in range(1, len(matches) + 1):
            markup.add(telebot.types.KeyboardButton(str(i)))

        markup.add(telebot.types.KeyboardButton("◀️ Назад"))

        msg = bot.send_message(chat_id, "❌ Введите номер из списка:", reply_markup=markup)
        bot.register_next_step_handler(msg, process_profession_choice)
        return

    matches = user_profession_matches.get(chat_id, [])

    if choice_num < 1 or choice_num > len(matches):
        markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)

        for i in range(1, len(matches) + 1):
            markup.add(telebot.types.KeyboardButton(str(i)))

        markup.add(telebot.types.KeyboardButton("◀️ Назад"))

        msg = bot.send_message(chat_id, f"❌ Введите номер от 1 до {len(matches)}:", reply_markup=markup)
        bot.register_next_step_handler(msg, process_profession_choice)
        return

    selected_profession = matches[choice_num - 1]
    user_selected_profession[chat_id] = selected_profession
    ask_for_additional_hazards(message, selected_profession)


def ask_for_additional_hazards(message, profession):
    """Спрашивает пользователя о дополнительных опасностях"""
    chat_id = message.chat.id
    user_states[chat_id] = "waiting_additional_hazards"
    user_selected_hazards[chat_id] = []

    ask_text = f"""✅ Выбрана профессия: {profession['name']} (№ {profession['code']})

❓ Есть ли еще дополнительные условия (идентифицированные опасности)?"""

    markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    markup.add(telebot.types.KeyboardButton("Да"), telebot.types.KeyboardButton("Нет"))
    markup.add(telebot.types.KeyboardButton("◀️ Назад"))

    msg = bot.send_message(chat_id, ask_text, reply_markup=markup)
    bot.register_next_step_handler(msg, process_hazard_question)


def process_hazard_question(message):
    """Обрабатывает ответ на вопрос об опасностях"""
    chat_id = message.chat.id

    if message.text == "◀️ Назад":
        user_states[chat_id] = "waiting_profession"
        user_selected_profession.pop(chat_id, None)
        user_selected_hazards.pop(chat_id, None)

        request_text = """🔍 Введи название или код профессии:"""

        markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
        markup.add(telebot.types.KeyboardButton("◀️ Назад"))

        msg = bot.send_message(chat_id, request_text, reply_markup=markup)
        bot.register_next_step_handler(msg, process_profession_input)
        return

    if message.text == "Да":
        user_states[chat_id] = "waiting_hazard_keyword"

        hazard_text = """🔍 Введи ключевое слово опасности:

Примеры:
• Шум
• Вибрация
• Электрический ток
• Пыль
• Высокая температура"""

        markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
        markup.add(telebot.types.KeyboardButton("Готово"))
        markup.add(telebot.types.KeyboardButton("◀️ Назад"))

        msg = bot.send_message(chat_id, hazard_text, reply_markup=markup)
        bot.register_next_step_handler(msg, process_hazard_keyword)
        return

    if message.text == "Нет":
        generate_final_document(message, chat_id)
        return

    markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    markup.add(telebot.types.KeyboardButton("Да"), telebot.types.KeyboardButton("Нет"))
    markup.add(telebot.types.KeyboardButton("◀️ Назад"))
    msg = bot.send_message(chat_id, "❌ Выберите 'Да' или 'Нет':", reply_markup=markup)
    bot.register_next_step_handler(msg, process_hazard_question)


def process_hazard_keyword(message):
    """Обрабатывает ввод ключевого слова опасности"""
    chat_id = message.chat.id

    if message.text == "◀️ Назад":
        profession = user_selected_profession.get(chat_id)
        ask_for_additional_hazards(message, profession)
        return

    if message.text == "Готово":
        generate_final_document(message, chat_id)
        return

    keyword = message.text.strip()

    if not keyword or len(keyword) < 2:
        markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
        markup.add(telebot.types.KeyboardButton("Готово"))
        markup.add(telebot.types.KeyboardButton("◀️ Назад"))
        msg = bot.send_message(chat_id, "❌ Введите минимум 2 символа\n\nПопробуйте еще раз:", reply_markup=markup)
        bot.register_next_step_handler(msg, process_hazard_keyword)
        return

    hazards_df = load_hazards_df()
    if hazards_df is None:
        bot.send_message(chat_id, "❌ Ошибка при загрузке файла опасностей")
        return

    matches = find_hazards_by_keyword(hazards_df, keyword)

    if not matches:
        markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
        markup.add(telebot.types.KeyboardButton("Готово"))
        markup.add(telebot.types.KeyboardButton("◀️ Назад"))
        msg = bot.send_message(chat_id, f"❌ Опасности по запросу '{keyword}' не найдены\n\nПопробуйте еще раз:",
                               reply_markup=markup)
        bot.register_next_step_handler(msg, process_hazard_keyword)
        return

    user_hazard_matches[chat_id] = matches
    user_states[chat_id] = "waiting_hazard_choice"

    message_text = f"🔍 Найдено опасностей по запросу '{keyword}':\n\n"

    markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)

    for i, match in enumerate(matches, 1):
        message_text += f"{i}. {match['name']}\n"
        markup.add(telebot.types.KeyboardButton(str(i)))

    message_text += "\n👇 Выбери номер нужной опасности (можно несколько раз выбирать разные):"

    markup.add(telebot.types.KeyboardButton("Готово"))
    markup.add(telebot.types.KeyboardButton("◀️ Назад"))
    msg = bot.send_message(chat_id, message_text, reply_markup=markup)
    bot.register_next_step_handler(msg, process_hazard_choice)


def process_hazard_choice(message):
    """Обрабатывает выбор опасности"""
    chat_id = message.chat.id

    if message.text == "◀️ Назад":
        user_states[chat_id] = "waiting_hazard_keyword"
        user_hazard_matches.pop(chat_id, None)

        hazard_text = """🔍 Введи ключевое слово опасности:"""

        markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
        markup.add(telebot.types.KeyboardButton("Готово"))
        markup.add(telebot.types.KeyboardButton("◀️ Назад"))

        msg = bot.send_message(chat_id, hazard_text, reply_markup=markup)
        bot.register_next_step_handler(msg, process_hazard_keyword)
        return

    if message.text == "Готово":
        # Проверяем, выбраны ли опасности, если нет - снова спрашиваем
        profession = user_selected_profession.get(chat_id)
        ask_for_additional_hazards(message, profession)
        return

    choice_text = message.text.strip()

    try:
        choice_num = int(choice_text)
    except ValueError:
        markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)

        matches = user_hazard_matches.get(chat_id, [])
        for i in range(1, len(matches) + 1):
            markup.add(telebot.types.KeyboardButton(str(i)))

        markup.add(telebot.types.KeyboardButton("Готово"))
        markup.add(telebot.types.KeyboardButton("◀️ Назад"))

        msg = bot.send_message(chat_id, "❌ Введите номер из списка:", reply_markup=markup)
        bot.register_next_step_handler(msg, process_hazard_choice)
        return

    matches = user_hazard_matches.get(chat_id, [])

    if choice_num < 1 or choice_num > len(matches):
        markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)

        for i in range(1, len(matches) + 1):
            markup.add(telebot.types.KeyboardButton(str(i)))

        markup.add(telebot.types.KeyboardButton("Готово"))
        markup.add(telebot.types.KeyboardButton("◀️ Назад"))

        msg = bot.send_message(chat_id, f"❌ Введите номер от 1 до {len(matches)}:", reply_markup=markup)
        bot.register_next_step_handler(msg, process_hazard_choice)
        return

    selected_hazard = matches[choice_num - 1]

    # Добавляем опасность в список выбранных
    if selected_hazard['name'] not in user_selected_hazards[chat_id]:
        user_selected_hazards[chat_id].append(selected_hazard['name'])

    # Показываем выбранные опасности
    selected_text = "✅ Выбранные опасности:\n"
    for hazard in user_selected_hazards[chat_id]:
        selected_text += f"• {hazard}\n"

    selected_text += "\n❓ Есть еще опасности?"

    markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    markup.add(telebot.types.KeyboardButton("Да"), telebot.types.KeyboardButton("Нет"))
    markup.add(telebot.types.KeyboardButton("◀️ Назад"))

    msg = bot.send_message(chat_id, selected_text, reply_markup=markup)
    bot.register_next_step_handler(msg, process_add_more_hazards)


def process_add_more_hazards(message):
    """Спрашивает, есть ли еще опасности"""
    chat_id = message.chat.id

    if message.text == "◀️ Назад":
        user_states[chat_id] = "waiting_hazard_keyword"
        user_hazard_matches.pop(chat_id, None)

        hazard_text = """🔍 Введи ключевое слово опасности:"""

        markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
        markup.add(telebot.types.KeyboardButton("Готово"))
        markup.add(telebot.types.KeyboardButton("◀️ Назад"))

        msg = bot.send_message(chat_id, hazard_text, reply_markup=markup)
        bot.register_next_step_handler(msg, process_hazard_keyword)
        return

    if message.text == "Да":
        user_states[chat_id] = "waiting_hazard_keyword"

        hazard_text = """🔍 Введи ключевое слово опасности:"""

        markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
        markup.add(telebot.types.KeyboardButton("Готово"))
        markup.add(telebot.types.KeyboardButton("◀️ Назад"))

        msg = bot.send_message(chat_id, hazard_text, reply_markup=markup)
        bot.register_next_step_handler(msg, process_hazard_keyword)
        return

    if message.text == "Нет":
        generate_final_document(message, chat_id)
        return

    markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    markup.add(telebot.types.KeyboardButton("Да"), telebot.types.KeyboardButton("Нет"))
    markup.add(telebot.types.KeyboardButton("◀️ Назад"))
    msg = bot.send_message(chat_id, "❌ Выберите 'Да' или 'Нет':", reply_markup=markup)
    bot.register_next_step_handler(msg, process_add_more_hazards)


def generate_final_document(message, chat_id):
    """Генерирует финальный документ"""
    profession = user_selected_profession.get(chat_id)
    hazards_list = user_selected_hazards.get(chat_id, [])

    if not profession:
        send_welcome(message)
        return

    processing_msg = bot.send_message(chat_id, f"⏳ Генерирую документ для профессии: {profession['name']}...")

    doc_output, error = generate_siz_document_from_match(profession, hazards_list if hazards_list else None)

    if error:
        bot.edit_message_text(error, chat_id, processing_msg.message_id)
        user_states.pop(chat_id, None)
        user_selected_profession.pop(chat_id, None)
        user_selected_hazards.pop(chat_id, None)
        return

    try:
        bot.delete_message(chat_id, processing_msg.message_id)
    except:
        pass

    doc_output.name = f"СИЗ_{profession['name']}.docx"

    caption = f"✅ Документ для профессии:\n{profession['name']} (№ {profession['code']})"
    if hazards_list:
        caption += f"\n\nДополнительные опасности:\n"
        for hazard in hazards_list:
            caption += f"• {hazard}\n"

    bot.send_document(chat_id, doc_output, caption=caption)

    user_states.pop(chat_id, None)
    user_selected_profession.pop(chat_id, None)
    user_selected_hazards.pop(chat_id, None)
    user_hazard_matches.pop(chat_id, None)

    markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    markup.add(telebot.types.KeyboardButton("📋 Расчет СИЗ"))
    markup.add(telebot.types.KeyboardButton("🏠 В главное меню"))

    bot.send_message(chat_id, "Готово! Что дальше?", reply_markup=markup)


@bot.message_handler(func=lambda message: message.text == "🏠 В главное меню")
def go_back_main(message):
    chat_id = message.chat.id
    user_states.pop(chat_id, None)
    user_profession_matches.pop(chat_id, None)
    user_selected_profession.pop(chat_id, None)
    user_selected_hazards.pop(chat_id, None)
    user_hazard_matches.pop(chat_id, None)
    send_welcome(message)


@bot.message_handler(func=lambda message: message.text == "📚 Терминология")
def terminology(message):
    bot.send_message(message.chat.id, "Функция терминологии в разработке...")


@bot.message_handler(func=lambda message: message.text == "⚖️ Законодательство")
def legislation(message):
    bot.send_message(message.chat.id, "Функция законодательства в разработке...")


@bot.message_handler(func=lambda message: True)
def echo_all(message):
    bot.send_message(message.chat.id, "Пожалуйста, используй кнопки меню для навигации или команду /start")


load_professions_df()
load_hazards_df()

bot.infinity_polling()