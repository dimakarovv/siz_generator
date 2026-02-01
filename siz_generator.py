import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os


def parse_norm_string(norm_str):
    """Парсит строку норм выдачи и извлекает количество и единицу измерения"""
    if pd.isna(norm_str):
        return None, None

    norm_str = str(norm_str).strip()

    # Если указано "до износа"
    if "износа" in norm_str.lower():
        return "до износа", None

    # Ищем количество в начале строки
    match = re.match(r'(\d+)\s*([а-яА-Я]+)', norm_str)
    if match:
        quantity = int(match.group(1))
        unit = match.group(2).strip()

        # Нормализуем единицы измерения
        unit_map = {
            'пара': 'пары',
            'пары': 'пары',
            'шт': 'штуки',
            'штука': 'штуки',
            'штуки': 'штуки',
            'комплект': 'комплекты',
            'комплекты': 'комплекты',
            'мл': 'мл'
        }

        normalized_unit = unit_map.get(unit.lower(), unit)
        return quantity, normalized_unit

    return None, None


def extract_years_info(norm_str):
    """Извлекает информацию о периоде (на 2 года, на 5 лет и т.д.)"""
    if pd.isna(norm_str):
        return ""

    norm_str = str(norm_str).lower()

    if "на 2 года" in norm_str:
        return " на 2 года"
    elif "на 3 года" in norm_str:
        return " на 3 года"
    elif "на 5 лет" in norm_str:
        return " на 5 лет"
    elif "на 1,5 года" in norm_str:
        return " на 1,5 года"

    return ""


def find_profession_by_code(professions_df, code):
    """Ищет профессию по коду (N п/п)"""
    try:
        code = int(code)
    except:
        return None, None, []

    matches = professions_df[professions_df.iloc[:, 0] == code]

    if matches.empty:
        return None, None, []

    start_idx = matches.index[0]
    profession_name = professions_df.iloc[start_idx, 1]

    return code, profession_name, start_idx


def find_professions_by_name(professions_df, search_name):
    """Ищет профессии по названию (может быть несколько совпадений)"""
    search_name_lower = search_name.lower().strip()

    matches = []

    for idx, row in professions_df.iterrows():
        # Проверяем наличие кода в первом столбце
        if pd.notna(row.iloc[0]) and isinstance(row.iloc[0], (int, float)):
            code = int(row.iloc[0])
            profession_name = row.iloc[1]

            if pd.notna(profession_name):
                profession_name_str = str(profession_name).strip()
                # Ищем частичное совпадение
                if search_name_lower in profession_name_str.lower():
                    matches.append({
                        'code': code,
                        'name': profession_name_str,
                        'idx': idx
                    })

    return matches


def select_profession(professions_df, search_query):
    """Выбирает профессию по коду или названию, уточняя при необходимости"""

    # Пытаемся найти по коду
    code, profession_name, start_idx = find_profession_by_code(professions_df, search_query)

    if code is not None:
        print(f"\nНайдена профессия по коду: {profession_name} (код: {code})")
        return code, profession_name, start_idx

    # Если не найдено по коду, ищем по названию
    print(f"\nПоиск профессии по названию: '{search_query}'")
    matches = find_professions_by_name(professions_df, search_query)

    if not matches:
        print(f"Профессии с названием, содержащим '{search_query}', не найдены")
        return None, None, None

    if len(matches) == 1:
        match = matches[0]
        print(f"Найдена профессия: {match['name']} (код: {match['code']})")
        return match['code'], match['name'], match['idx']

    # Если найдено несколько совпадений, уточняем выбор
    print(f"\nНайдено несколько профессий, соответствующих вашему запросу:")
    print("Какую из данных профессий вам нужна?")

    for i, match in enumerate(matches, 1):
        print(f"{i}. {match['name']} (№ п/п {match['code']})")

    while True:
        try:
            choice = int(input(f"\nВведите номер профессии (1-{len(matches)}): "))
            if 1 <= choice <= len(matches):
                selected = matches[choice - 1]
                print(f"\nВы выбрали: {selected['name']} (код: {selected['code']})")
                return selected['code'], selected['name'], selected['idx']
            else:
                print(f"Пожалуйста, введите число от 1 до {len(matches)}")
        except ValueError:
            print("Ошибка: введите корректное число")


def extract_siz_data(professions_df, profession_code, start_idx):
    """Извлекает данные о СИЗ для конкретной профессии"""

    profession_name = professions_df.iloc[start_idx, 1]

    print(f"\nИзвлечение данных СИЗ для профессии: {profession_name}")

    siz_list = []

    # Обрабатываем строки, начиная со следующей после названия профессии
    for idx in range(start_idx + 1, len(professions_df)):
        row = professions_df.iloc[idx]

        # Проверяем, не является ли это новой профессией (новый код в первом столбце)
        if pd.notna(row.iloc[0]) and isinstance(row.iloc[0], (int, float)) and not pd.isna(row.iloc[0]):
            # Это новая профессия, прерываем цикл
            break

        # Получаем наименование СИЗ (четвёртый столбец, индекс 3)
        siz_name = row.iloc[3] if len(row) > 3 else None

        # Получаем нормы выдачи (пятый столбец, индекс 4)
        norm_str = row.iloc[4] if len(row) > 4 else None

        if pd.notna(siz_name) and str(siz_name).strip():
            quantity, unit = parse_norm_string(norm_str)
            years_info = extract_years_info(norm_str)

            # Формируем строку количества на год
            if quantity:
                quantity_str = f"{quantity}{years_info}"
            else:
                quantity_str = str(norm_str) if norm_str else "до износа"

            siz_list.append({
                'name': str(siz_name).strip(),
                'unit': unit if unit else 'штуки',
                'quantity': quantity_str,
                'norm': '№767н'
            })

    return siz_list


def find_siz_table(doc):
    """Находит таблицу СИЗ в документе Word"""
    for idx, table in enumerate(doc.tables):
        if len(table.rows) > 0:
            first_cell = table.rows[0].cells[0].text.strip()
            # Ищем таблицу, которая начинается с "Наименование СИЗ"
            if "Наименование СИЗ" in first_cell:
                print(f"Найдена таблица СИЗ с индексом {idx}")
                print(f"Количество строк в таблице: {len(table.rows)}")
                print(f"Количество столбцов в таблице: {len(table.columns)}")
                return table, idx

    return None, None


def set_cell_border(cell, **kwargs):
    """Устанавливает границы для ячейки таблицы"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # Создаем элемент tcBorders если его нет
    tcBorders = OxmlElement('w:tcBorders')

    for edge in ('top', 'left', 'bottom', 'right'):
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), 'single')
        edge_el.set(qn('w:sz'), '12')
        edge_el.set(qn('w:space'), '0')
        edge_el.set(qn('w:color'), '000000')
        tcBorders.append(edge_el)

    tcPr.append(tcBorders)


def add_table_row(table, siz_data):
    """Добавляет новую строку в таблицу с правильным форматированием"""
    # Добавляем новую строку
    new_row = table.add_row()

    # Копируем стиль из последней заполненной строки если она есть
    if len(table.rows) > 2:
        source_row = table.rows[-2]
        for source_cell, target_cell in zip(source_row.cells, new_row.cells):
            # Копируем границы
            set_cell_border(target_cell)
    else:
        # Применяем стандартные границы
        for cell in new_row.cells:
            set_cell_border(cell)

    return new_row


def fill_siz_table(doc, siz_data):
    """Заполняет таблицу СИЗ данными"""

    table, table_idx = find_siz_table(doc)

    if not table:
        print("Таблица СИЗ не найдена в документе!")
        return False

    print(f"Найдено {len(siz_data)} видов СИЗ\n")

    # Определяем количество строк для данных (пропускаем заголовок)
    # Обычно заголовок находится в первой строке
    header_rows = 1
    existing_data_rows = len(table.rows) - header_rows
    required_rows = len(siz_data)

    print(f"Существующих строк для данных: {existing_data_rows}")
    print(f"Требуется строк: {required_rows}")

    # Если нужно больше строк, добавляем их
    if required_rows > existing_data_rows:
        rows_to_add = required_rows - existing_data_rows
        print(f"Добавляем {rows_to_add} новых строк\n")
        for i in range(rows_to_add):
            add_table_row(table, siz_data)

    # Заполняем данные
    for idx, siz in enumerate(siz_data):
        row_idx = header_rows + idx
        row = table.rows[row_idx]

        print(f"Строка {row_idx + 1}: {siz['name']}")

        # Заполняем ячейки в строке
        if len(row.cells) >= 4:
            # Очищаем ячейки
            for cell in row.cells[:4]:
                cell.text = ""

            # Заполняем данные
            row.cells[0].text = siz['name']
            row.cells[1].text = siz['norm']
            row.cells[2].text = siz['unit']
            row.cells[3].text = siz['quantity']
        else:
            print(f"  Внимание: недостаточно столбцов в таблице (требуется 4, есть {len(row.cells)})")

    return True


def fill_personal_anketa_docx(template_path, profession_query, output_path):
    """Заполняет личную карточку СИЗ в Word файле на основе данных профессии"""

    # Читаем файл профессий
    professions_df = pd.read_excel('professions.xlsx', header=None)

    print(f"Общее количество строк в файле профессий: {len(professions_df)}")
    print(f"Общее количество столбцов в файле профессий: {len(professions_df.columns)}")

    # Выбираем профессию (по коду или названию)
    profession_code, profession_name, start_idx = select_profession(professions_df, profession_query)

    if profession_code is None:
        print("\nНе удалось найти подходящую профессию")
        return False

    # Извлекаем данные СИЗ
    siz_data = extract_siz_data(professions_df, profession_code, start_idx)

    if not siz_data:
        print(f"Не удалось найти СИЗ для профессии '{profession_name}' (код {profession_code})")
        return False

    # Загружаем шаблон Word документа
    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"Ошибка при открытии документа: {e}")
        return False

    print(f"\nДокумент загружен, количество таблиц: {len(doc.tables)}\n")

    # Заполняем таблицу СИЗ
    if not fill_siz_table(doc, siz_data):
        return False

    # Формируем имя выходного файла
    safe_profession_name = re.sub(r'[<>:"/\\|?*]', '', str(profession_name))
    output_filename = f"{safe_profession_name}_{profession_code}.docx"

    if output_path:
        output_filename = os.path.join(output_path, output_filename)

    # Сохраняем файл
    try:
        doc.save(output_filename)
        print(f"\nФайл успешно создан: {output_filename}")
        return True
    except Exception as e:
        print(f"Ошибка при сохранении документа: {e}")
        return False


# Основная программа
if __name__ == "__main__":
    # Пример использования
    print("=" * 60)
    print("Генератор личных карточек учета выдачи СИЗ")
    print("=" * 60)

    profession_query = input("\nВведите код профессии (N п/п) или название профессии: ").strip()

    if not profession_query:
        print("Ошибка: введите код или название профессии")
    else:
        fill_personal_anketa_docx('personal_anketa.docx', profession_query, '.')